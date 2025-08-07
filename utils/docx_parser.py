from docx import Document
import docx.oxml.shared
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import os

# --- Fungsi bungkus teks dengan tanda hubung ---
def wrap_text_with_hyphenation(text, max_line_length=52):
    lines = []
    for original_line in text.split('\n'):
        words = original_line.split()
        current_line = ""
        for word in words:
            while len(word) > max_line_length:
                split_point = max_line_length - 1
                lines.append(word[:split_point] + "-")
                word = word[split_point:]
            if len(current_line) + len(word) + (1 if current_line else 0) > max_line_length:
                if current_line:
                    lines.append(current_line)
                current_line = word
            else:
                current_line += (" " if current_line else "") + word
        if current_line:
            lines.append(current_line)
    return '\n'.join(lines)

# --- Fungsi pengganti teks di paragraf dan tabel ---
def replace_all_placeholders(doc, replacements):
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                inline_replace(p, key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in p.text:
                            inline_replace(p, key, val)

def inline_replace(paragraph, key, val):
    text = paragraph.text.replace(key, val)
    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph._element.remove(paragraph.runs[i]._element)
    paragraph.add_run(text)

# --- Fungsi sisip paragraf bernomor ---
def insert_numbered_paragraphs_in_tables(doc, placeholder, items, style):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for i, para in enumerate(cell.paragraphs):
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, "")
                            p_element = para._element
                            for item in items:
                                paragraph = cell.add_paragraph(style=style)
                                parts = item.split('\n')
                                for idx, part in enumerate(parts):
                                    if idx > 0:
                                        br = OxmlElement("w:br")
                                        paragraph._element.append(br)
                                    paragraph.add_run(part)
                            p_element.getparent().remove(p_element)
                            return

# --- Fungsi keepLines dan keepNext ---
def set_keep_with_next(paragraph, keep=True):
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    if keep:
        pPr.append(keepNext)

def set_keep_lines(paragraph, keep=True):
    pPr = paragraph._element.get_or_add_pPr()
    keepLines = OxmlElement('w:keepLines')
    if keep:
        pPr.append(keepLines)

def keep_paragraphs_together(doc, start_text):
    found = False
    for p in doc.paragraphs:
        if start_text in p.text:
            found = True
        if found:
            set_keep_with_next(p)
            set_keep_lines(p)

# === FUNGSI UTAMA ===
def generate_surat_tugas(peserta, surat_info, no_template=2, template_path=None, output_path=None, header_row_count=2):
    if template_path is None:
        template_path = "templates/surat_tugas_template - Nama - NIP Combined.docx" if no_template == 2 else "templates/surat_tugas_template.docx"
    if output_path is None:
        output_path = "output.docx"

    doc = Document(template_path)

    # Placeholder biasa
    replacements = {
        "{{nomor_surat_tugas}}": surat_info.get("nomor_surat_tugas", ""),
        "{{NAMA_KEGIATAN}}": surat_info.get("nama_kegiatan", ""),
        "{{TAHUN_ANGGARAN_KEGIATAN}}": surat_info.get("tahun_anggaran_kegiatan", ""),
        "{{HARI_PELAKSANAAN}}": surat_info.get("hari_pelaksanaan", ""),
        "{{TANGGAL_PELAKSANAAN}}": surat_info.get("tanggal_pelaksanaan", ""),
        "{{TEMPAT_PELAKSANAAN}}": surat_info.get("tempat_pelaksanaan", ""),
        "{{KOTA}}": surat_info.get("kota", ""),
        "{{TANGGAL}}": surat_info.get("tanggal", ""),
    }

    replace_all_placeholders(doc, replacements)

    # List menimbang & dasar hukum
    menimbang = surat_info.get("menimbang", [])
    if isinstance(menimbang, str):
        menimbang = [menimbang]
    dasar_hukum = surat_info.get("dasar_hukum", [])
    if isinstance(dasar_hukum, str):
        dasar_hukum = [dasar_hukum]

    menimbang_wrapped = [wrap_text_with_hyphenation(item, 52) for item in menimbang]
    dasar_hukum_wrapped = [wrap_text_with_hyphenation(item, 52) for item in dasar_hukum]

    if len(menimbang) == 1:
        insert_numbered_paragraphs_in_tables(doc, "{{menimbang}}", menimbang_wrapped, style='Normal Font')
    else:
        insert_numbered_paragraphs_in_tables(doc, "{{menimbang}}", menimbang_wrapped, style='List menimbangs')

    if len(dasar_hukum) == 1:
        insert_numbered_paragraphs_in_tables(doc, "{{dasar_hukum}}", dasar_hukum_wrapped, style='Normal Font')
    else:
        insert_numbered_paragraphs_in_tables(doc, "{{dasar_hukum}}", dasar_hukum_wrapped, style='List numberg')

    # === ISI TABEL PESERTA ===
    try:
        table = doc.tables[1]
    except IndexError:
        raise IndexError("Template tidak memiliki tabel kedua (index 1). Periksa template.")

    if header_row_count < 1:
        header_row_count = 1
    while len(table.rows) > header_row_count:
        table._tbl.remove(table.rows[header_row_count]._tr)

    for idx, item in enumerate(peserta, start=1):
        row_cells = table.add_row().cells
        if no_template == 1:
            values = [f"{idx}.", item.nama, item.nip, item.jabatan, item.satker]
        else:
            if len(item.nip) == 18:
                formatted_nip = f"{item.nip[:8]} {item.nip[8:14]} {item.nip[14:15]} {item.nip[15:]}"
            else:
                formatted_nip = item.nip
            nama_nip = f"{item.nama}\nNIP. {formatted_nip}"
            values = [f"{idx}.", nama_nip, item.jabatan, item.satker]

        for i in range(min(len(row_cells), len(values))):
            cell = row_cells[i]
            cell_para = cell.paragraphs[0]
            for r in cell_para.runs:
                cell_para._element.remove(r._element)
            cell_para.add_run(values[i])

            # === Penentuan alignment berdasarkan kolom ===
            try:
                if i == 0:
                    # Kolom NO → align center
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif no_template == 2 and i in [2, 3]:
                    # Kolom jabatan dan satker → align left
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif no_template == 1 and i in [3, 4]:
                    # Untuk no_template == 1, jabatan & satker ada di kolom 3 dan 4
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                else:
                    # Default (nama, nip, dll)
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass

    # === Header Tabel ===
    if len(table.rows) >= header_row_count:
        for header_row_idx in range(header_row_count):
            tr = table.rows[header_row_idx]._tr
            trPr_list = tr.xpath('./w:trPr')
            trPr_element = trPr_list[0] if trPr_list else OxmlElement('w:trPr')
            if not trPr_list:
                tr.append(trPr_element)
            if not trPr_element.xpath('./w:tblHeader'):
                trPr_element.append(docx.oxml.shared.OxmlElement('w:tblHeader'))

            for i, cell in enumerate(table.rows[header_row_idx].cells):
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    keep_paragraphs_together(doc, "Untuk")

    doc.save(output_path)
    return output_path
